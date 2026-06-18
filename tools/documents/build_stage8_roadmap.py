from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
OUT = ROOT / "docs" / "reports"
DOCX = OUT / "Garmetix-Stage8-Development-Roadmap-v3.12.0.docx"

NAVY = "102A43"
TEAL = "00A6A6"
CYAN = "DFF7F6"
BLUE_GRAY = "E8EEF5"
LIGHT = "F4F6F9"
MID = "52606D"
DARK = "1F2933"
WHITE = "FFFFFF"
GREEN = "137333"
AMBER = "8A5A00"
RED = "A61B1B"


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_widths(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    total = sum(widths)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(widths[idx] / 1440)
            set_cell_margins(cell)


def set_run(run, size=10.5, bold=False, color=DARK, font="Aptos"):
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def add_text(doc, text, style=None, size=10.5, bold=False, color=DARK, after=6, align=None):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.16
    set_run(p.add_run(text), size=size, bold=bold, color=color)
    return p


def add_stage_banner(doc, code, title, version, source):
    table = doc.add_table(rows=1, cols=2)
    set_table_widths(table, [1750, 7610])
    left, right = table.rows[0].cells
    shade(left, TEAL)
    shade(right, NAVY)
    left.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    right.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = left.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    set_run(p.add_run(code), size=15, bold=True, color=WHITE)
    p = right.paragraphs[0]
    p.paragraph_format.space_after = Pt(1)
    set_run(p.add_run(title), size=13.5, bold=True, color=WHITE)
    p = right.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    set_run(p.add_run(f"{version} | {source}"), size=8.5, color="D7E3EF")
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_work_table(doc, rows):
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    set_table_widths(table, [2250, 4650, 2460])
    headers = ["Workstream", "Implementation details", "Completion evidence"]
    for idx, text in enumerate(headers):
        cell = table.rows[0].cells[idx]
        shade(cell, BLUE_GRAY)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(0)
        set_run(p.add_run(text), size=9, bold=True, color=NAVY)
    set_repeat_table_header(table.rows[0])
    for workstream, details, evidence in rows:
        cells = table.add_row().cells
        for idx, text in enumerate((workstream, details, evidence)):
            cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cells[idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            set_run(p.add_run(text), size=8.8, bold=(idx == 0), color=NAVY if idx == 0 else DARK)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_callout(doc, label, text, fill=CYAN, label_color=TEAL):
    table = doc.add_table(rows=1, cols=1)
    set_table_widths(table, [9360])
    cell = table.cell(0, 0)
    shade(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    set_run(p.add_run(f"{label}: "), size=9.5, bold=True, color=label_color)
    set_run(p.add_run(text), size=9.5, color=DARK)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def add_footer(section):
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    set_run(p.add_run("Garmetix Stage 8 Development Roadmap  |  "), size=8, color=MID)
    run = p.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr, fld_char2])


def add_header(section):
    header = section.header
    p = header.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    set_run(p.add_run("GARMETIX"), size=8.5, bold=True, color=TEAL)
    set_run(p.add_run("  |  Product and Engineering Plan"), size=8.5, color=MID)


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.78)
section.bottom_margin = Inches(0.72)
section.left_margin = Inches(0.85)
section.right_margin = Inches(0.85)
section.header_distance = Inches(0.35)
section.footer_distance = Inches(0.35)
add_header(section)
add_footer(section)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Aptos"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Aptos")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos")
normal.font.size = Pt(10.5)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.16

for name, size, color, before, after in (
    ("Title", 27, NAVY, 0, 7),
    ("Subtitle", 13, MID, 0, 12),
    ("Heading 1", 17, NAVY, 15, 7),
    ("Heading 2", 13, TEAL, 11, 5),
    ("Heading 3", 11, NAVY, 8, 4),
):
    style = styles[name]
    style.font.name = "Aptos Display" if name in ("Title", "Heading 1", "Heading 2") else "Aptos"
    style._element.rPr.rFonts.set(qn("w:ascii"), style.font.name)
    style._element.rPr.rFonts.set(qn("w:hAnsi"), style.font.name)
    style.font.size = Pt(size)
    style.font.bold = name != "Subtitle"
    style.font.color.rgb = RGBColor.from_string(color)
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.keep_with_next = True

# Cover
add_text(doc, "GARMENT STORE MANAGEMENT PLATFORM", size=9, bold=True, color=TEAL, after=18)
add_text(doc, "Garmetix", style="Title", size=31, bold=True, color=NAVY, after=3)
add_text(doc, "Stage 8 Development Roadmap", style="Subtitle", size=17, color=MID, after=8)
add_text(
    doc,
    "From Stage 7M / Version 3.12.0 to production-ready v4.x",
    size=11.5,
    bold=True,
    color=TEAL,
    after=24,
)

cover = doc.add_table(rows=5, cols=2)
cover.style = "Table Grid"
set_table_widths(cover, [2100, 7260])
cover_data = [
    ("Current baseline", "Stage 7M - Pre-v4.0 UI Naming and Menu Cleanup"),
    ("Version / build", "3.12.0 / GARMETIX-7M-20260611-3120"),
    ("Branch / commit", "Version3.0 / 6d440a5"),
    ("Prepared", "11 June 2026"),
    ("Roadmap scope", "Baseline stabilization, UI, access, returns, inventory, accounting, tests, integrations, and go-live"),
]
for row, (label, value) in zip(cover.rows, cover_data):
    shade(row.cells[0], BLUE_GRAY)
    for idx, text in enumerate((label, value)):
        row.cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = row.cells[idx].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        set_run(p.add_run(text), size=9.5, bold=(idx == 0), color=NAVY if idx == 0 else DARK)

doc.add_paragraph().paragraph_format.space_after = Pt(14)
add_callout(
    doc,
    "Purpose",
    "Define the implementation sequence, boundaries, acceptance evidence, and production gates for the next development cycle while preserving the Stage 7M user experience and accounting architecture.",
)
add_callout(
    doc,
    "Source note",
    "Stage 8A and Stage 8B incorporate the Stage 8 direction preserved locally from the web-chat planning history. The private 'Garmetix Web 2' conversation itself was not available as an exported or public transcript, so this report does not invent or quote unseen chat content. Stage 8C through Stage 8G consolidate the current repository audit and remaining engineering gaps.",
    fill=LIGHT,
    label_color=AMBER,
)

add_page_break(doc)

# Executive summary
add_text(doc, "Executive Summary", style="Heading 1")
add_text(
    doc,
    "Garmetix already contains broad operational coverage across billing, purchase, inventory, accounting, banking, GST, HR, payroll, access control, backup, reporting, audit, and Oracle secondary synchronization. Stage 8 is therefore a hardening and completion program rather than a restart.",
)
add_text(
    doc,
    "The implementation sequence begins with known baseline defects, then standardizes the user experience, strengthens access administration, formalizes purchase-return and inventory documents, improves accounting fidelity, introduces dependable automated testing, and finally validates external integrations and deployment.",
)

summary = doc.add_table(rows=1, cols=4)
summary.style = "Table Grid"
set_table_widths(summary, [1280, 2520, 3160, 2400])
for idx, text in enumerate(("Phase", "Primary outcome", "Business value", "Release gate")):
    shade(summary.rows[0].cells[idx], NAVY)
    p = summary.rows[0].cells[idx].paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    set_run(p.add_run(text), size=8.8, bold=True, color=WHITE)
set_repeat_table_header(summary.rows[0])
summary_rows = [
    ("P0", "Stable baseline", "Reliable dashboard, clean warnings, repeatable validation", "No known response-body defect"),
    ("8A", "Consistent v4.0 UI", "Faster daily work and fewer layout errors", "All UI Audit pages reviewed"),
    ("8B", "Permission confidence", "Correct role boundaries and administration", "Role matrix passes"),
    ("8C", "Formal purchase returns", "Traceable stock, GST, vendor and ledger settlement", "Return lifecycle reconciles"),
    ("8D", "Authoritative inventory", "Defensible stock and valuation", "Movement ledger reconciles"),
    ("8E", "Accounting fidelity", "Accurate split payments, banking and due balances", "Journals and subledgers balance"),
    ("8F", "Automated assurance", "Safer releases and fewer regressions", "Critical suites pass"),
    ("8G", "Production go-live", "Validated integrations, backup and deployment", "Release checklist signed off"),
]
for row_data in summary_rows:
    cells = summary.add_row().cells
    for idx, text in enumerate(row_data):
        if idx == 0:
            shade(cells[idx], CYAN)
        p = cells[idx].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        set_run(p.add_run(text), size=8.2, bold=(idx == 0), color=NAVY if idx == 0 else DARK)

add_text(doc, "Delivery Principles", style="Heading 2")
principles = [
    ("Preserve business boundaries", "Regular accounting vouchers and Off Book Cash Vouchers remain separate."),
    ("Prefer formal documents", "Transactions that affect stock, GST, vendor/customer balances, or audit must have traceable header/detail records."),
    ("Keep background flags internal", "Party-ledger and bank-ledger linkage flags are derived by the system and are not user-entered controls."),
    ("Build for daily operation", "Large master-detail workflows use full pages or wide workspaces, not narrow side drawers."),
    ("Prove completion", "Every stage ends with build, runtime, reconciliation, permission, and visual evidence."),
]
for title, body in principles:
    add_callout(doc, title, body, fill=LIGHT, label_color=TEAL)

add_page_break(doc)

# Baseline and dependencies
add_text(doc, "1. Baseline and Immediate Stabilization", style="Heading 1")
add_callout(
    doc,
    "Stage 7M contract",
    "Preserve the grouped Nuxt UI shell, menu naming, current routes, permission-aware navigation, dark theme default, Off Book grouping, contextual Store/Company dashboards, and NUXT_PUBLIC_DASHBOARD_SHELL=legacy rollback.",
)
add_work_table(
    doc,
    [
        ("Dashboard route", "Correct the Minimal API mapping so GET /api/dashboard/home writes DashboardHomeDto to the response.", "Authenticated request returns valid JSON and smart redirect works."),
        ("Compiler quality", "Resolve nullable warnings in purchase receipt mapping and data-consistency number handling.", "Release build completes without project warnings."),
        ("Build reliability", "Eliminate external font metadata dependency or configure a reliable system-CA/local-font path.", "Nuxt build succeeds without provider certificate failures."),
        ("Version identity", "Keep API, frontend package, visible release, stage and build code synchronized.", "System Info reports a full match."),
        ("Validation entry point", "Create one current aggregate validation command for backend, frontend, static checks and smoke tests.", "Single command produces a clear pass/fail report."),
        ("Docker runtime", "Verify login, dashboard, workspace selection, role navigation, API health and database readiness.", "Documented smoke run passes on the current Compose stack."),
    ],
)
add_text(doc, "Exit gate", style="Heading 2")
add_text(
    doc,
    "Stage 8 feature work begins only after the dashboard response defect is fixed, builds are repeatable, version identity matches, and the current Docker stack passes a basic authenticated smoke test.",
)

add_page_break(doc)

# Stage 8A
add_stage_banner(doc, "8A", "Full UI Audit and Standardization", "Target v4.0", "Web Chat Stage 8 direction preserved locally")
add_text(
    doc,
    "Objective: make every operational page feel like one coherent Nuxt UI application, with practical forms and predictable navigation for repeated store work.",
)
add_work_table(
    doc,
    [
        ("Page framework", "Standardize page title, context, primary action, filters, toolbar, table, pagination, loading, error and empty states.", "Every route follows the approved page pattern."),
        ("Large forms", "Move invoice, purchase, voucher, employee, payroll and other master-detail forms to dedicated pages or wide workspaces.", "No critical workflow is constrained by a narrow drawer."),
        ("Responsive shell", "Verify mobile drawer, collapsed icons, active groups, command palette, footer menus, workspace selector and topbar actions.", "No overlap or inaccessible control at mobile, tablet or desktop sizes."),
        ("Tables and actions", "Contain horizontal overflow, keep actions stable, wrap buttons cleanly and preserve readable columns.", "Tables never break the shell or hide essential actions."),
        ("Notification UX", "Add a useful action/notification surface for operational exceptions without exposing technical status noise.", "Users can review and act on meaningful alerts."),
        ("Reusable components", "Extract shared page-header, filter, form-section, master-detail and status patterns where duplication is material.", "Reduced repeated markup without obscuring module logic."),
        ("Visible language", "Remove stale stage/migration language and normalize capitalization, labels, help text and errors.", "Business pages contain only user-relevant language."),
        ("Visual QA", "Use /ui-audit plus Playwright screenshots at representative desktop and mobile viewports.", "Every audit item is reviewed and recorded."),
    ],
)
add_text(doc, "Module order", style="Heading 2")
add_text(
    doc,
    "Sales and Billing; Purchase and Returns; Inventory and Product Master; Accounting and Banking; HR and Payroll; Reports and GST; CRM; Admin, Data, Maintenance and System.",
)

add_page_break(doc)

# Stage 8B
add_stage_banner(doc, "8B", "User, Role and Permission Hardening", "Target v4.1", "Web Chat Stage 8 direction preserved locally")
add_text(
    doc,
    "Objective: turn the existing access screens and policies into a verified administration system with clear scope, action rights and audit evidence.",
)
add_work_table(
    doc,
    [
        ("User lifecycle", "Audit create, edit, activate/deactivate, password reset and assigned workspace flows.", "Admin can manage users without direct database work."),
        ("Role assignment", "Validate Owner, Admin, PowerUser, Accountant, Remote Accountant, StoreManager, Salesman, HR and Payroll behavior.", "Each role sees only approved menus and data."),
        ("Action rights", "Enforce view, entry, edit and delete separately in both API and frontend.", "StoreManager has view/entry but restricted edit/delete; delete remains Owner/Admin."),
        ("Admin isolation", "Keep Company, Roles/Users, onboarding, seeding and administrative tools hidden and blocked for unauthorized roles.", "Direct URL and API attempts are denied."),
        ("Scope rules", "Validate company, store-group and store locking, defaults and workspace changes.", "Transactions cannot escape assigned scope."),
        ("Audit trail", "Record role, scope, status and password-administration changes.", "Each administrative change has actor, time and target."),
        ("Automated matrix", "Create integration and browser tests for menu, route, API and action permissions.", "Permission suite passes for every supported role."),
    ],
)
add_callout(
    doc,
    "Important clarification",
    "The user list and core access UI already exist. Stage 8B is an audit, refinement and test phase, not a claim that user administration must be created from nothing.",
    fill=LIGHT,
)

add_page_break(doc)

# Stage 8C
add_stage_banner(doc, "8C", "Purchase Return and Vendor Settlement", "Post-v4.1 business completion", "Repository gap analysis")
add_text(
    doc,
    "Objective: replace inferred purchase-return behavior with a formal, auditable document lifecycle tied to stock, GST, vendor balances and settlement.",
)
add_work_table(
    doc,
    [
        ("Data model", "Add PurchaseReturn and PurchaseReturnItem with original purchase references and immutable item/tax snapshots.", "Return remains understandable even if masters change later."),
        ("Return workflow", "Support partial quantities, reasons, validation against available return quantity, approval/cancellation and history.", "No item can be over-returned."),
        ("Stock posting", "Create linked outward stock movements and reverse them on controlled cancellation.", "Return document and movement ledger reconcile exactly."),
        ("GST treatment", "Calculate item-level HSN, CGST, SGST and IGST reversal for input tax credit.", "GST return data matches return items."),
        ("Commercial note", "Create/link vendor debit notes and provide A4/A5/PDF print actions.", "Return, note and original purchase are mutually traceable."),
        ("Settlement", "Support vendor refund receipt, balance adjustment, replacement and outstanding carry-forward.", "Vendor subledger equals the chosen settlement path."),
        ("Accounting", "Post balanced journals through the existing accounting service.", "Journal, vendor ledger and bank/cash account reconcile."),
        ("Audit", "Record create, approve, settle, cancel and reprint events.", "Complete lifecycle visible from the return detail page."),
    ],
)

add_page_break(doc)

# Stage 8D
add_stage_banner(doc, "8D", "Inventory Documents and Valuation", "Inventory authority program", "Repository gap analysis")
add_text(
    doc,
    "Objective: make inventory defensible through formal documents, an authoritative movement ledger, consistent valuation and reconciliation reports.",
)
add_work_table(
    doc,
    [
        ("Adjustment document", "Add header/items for damage, shortage, excess, write-off and correction with reason and approval.", "Each adjustment has a source document."),
        ("Transfer document", "Add source/destination stores, dispatch, receive, variance and in-transit states.", "Both stores reconcile without duplicate stock."),
        ("Physical count", "Add count sessions, snapshots, counted quantities, variances, approval and posting.", "Posted variance links to movement and journal entries."),
        ("Movement authority", "Derive stock balance from StockMovement rather than cumulative purchase/sold counters.", "Movement sum equals displayed and report stock."),
        ("Valuation", "Implement weighted-average cost first and document treatment of returns, transfers and adjustments.", "Valuation report reproduces ledger value."),
        ("Accounting", "Post shortage, excess, write-off and transfer implications where required.", "Inventory value and accounting remain aligned."),
        ("Analytics", "Add ageing, low-stock risk, slow-moving, valuation and reconciliation reports.", "Reports support store and company scopes."),
        ("Concurrency", "Test sequence safety, stock locks, negative-stock policy and competing transactions.", "Concurrent operations cannot oversell or corrupt stock."),
    ],
)

add_page_break(doc)

# Stage 8E
add_stage_banner(doc, "8E", "Accounting, Banking and Payment Hardening", "Financial fidelity program", "Repository gap analysis")
add_text(
    doc,
    "Objective: ensure every payment instrument, allocation and derived ledger relationship is complete, balanced and auditable under Indian accounting practice.",
)
add_work_table(
    doc,
    [
        ("Split payment posting", "Post one accounting line per cash, card, UPI, bank, cheque, advance, credit note, loyalty or store-credit row.", "Invoice payment rows reconcile to journal lines."),
        ("Instrument details", "Persist bank account, card/UPI reference, cheque number/date, issuer, depositor, beneficiary and purpose.", "Audit can identify who paid whom, how and why."),
        ("Value application", "Prevent duplicate use of advances, credit notes, loyalty and store credit.", "Each source has an application ledger and remaining balance."),
        ("Bank reconciliation", "Import/enter statement rows and match, partially match, unmatch or adjust transactions.", "Book balance and statement balance reconcile."),
        ("Cheque lifecycle", "Track issued, received, deposited, cleared, bounced, cancelled and replaced states.", "Cheque log and bank ledger remain synchronized."),
        ("Year controls", "Add financial-year lock, reopening authority and posting-date validation.", "Locked periods reject unauthorized changes."),
        ("Journal control", "Strengthen debit/credit balancing, reversal and duplicate-reference rules.", "No unbalanced journal can be posted."),
        ("Derived ledgers", "Keep Party and Bank Account ledger creation/linking internal and idempotent.", "Users never edit internal flags; links survive rename/edit."),
        ("Management views", "Add customer/vendor dues, cash/payment summary and store-group comparisons.", "Company dashboard exposes actionable financial positions."),
    ],
)

add_page_break(doc)

# Stage 8F/G
add_stage_banner(doc, "8F", "Deep Audit and Automated Testing", "Release assurance", "Repository gap analysis")
add_work_table(
    doc,
    [
        ("Change history", "Store before/after JSON for material entities with actor, reason, time and correlation reference.", "Auditors can reconstruct important changes."),
        ("Backend tests", "Add xUnit unit and API integration tests for domain rules and endpoints.", "Critical modules have repeatable automated coverage."),
        ("Database tests", "Exercise PostgreSQL transactions, locks, sequences, migrations and rollback behavior.", "Concurrency and migration scenarios pass."),
        ("Financial tests", "Cover journal balance, GST, returns, payroll, advances, dues and reconciliation.", "Expected ledgers match golden test cases."),
        ("Frontend tests", "Add Playwright flows for login, billing, purchase, returns, stock, payroll and administration.", "Critical user journeys pass in CI/local validation."),
        ("Recovery tests", "Restore a backup into a disposable database and verify manifests, checksums and key counts.", "Recovery drill has recorded evidence."),
    ],
)

add_stage_banner(doc, "8G", "External Integration and Production Go-Live", "Production release", "Repository gap analysis")
add_work_table(
    doc,
    [
        ("Email", "Validate SMTP reset and payslip delivery with production sender/domain.", "Delivered messages and failures are logged."),
        ("Google Drive", "Validate service account, folder permissions, upload, download and restore.", "Cloud backup and recovery drill passes."),
        ("GST provider", "Configure licensed GSTIN lookup credentials and validate production response mapping.", "Verified party data and mismatch alerts work."),
        ("Oracle Cloud", "Validate wallet/TNS, push/pull, review/apply, dead letters and one real external app.", "Ownership and final auto-apply policy are signed off."),
        ("Clean installation", "Run fresh Docker installation, migrations, onboarding, seed and first login.", "New installation reaches healthy operational state."),
        ("Production security", "Configure HTTPS, secrets, CORS, retention, scheduled backups and log rotation.", "Preflight reports no critical production issue."),
        ("Acceptance", "Run role matrix, module smoke tests, reports, printing, backup/restore and accounting reconciliation.", "Production release checklist is approved."),
    ],
)

add_page_break(doc)

# Sequencing and governance
add_text(doc, "2. Delivery Sequence and Governance", style="Heading 1")
sequence = doc.add_table(rows=1, cols=5)
sequence.style = "Table Grid"
set_table_widths(sequence, [1050, 1650, 2500, 2310, 1850])
for idx, text in enumerate(("Order", "Package", "Depends on", "Primary verification", "Decision gate")):
    shade(sequence.rows[0].cells[idx], NAVY)
    p = sequence.rows[0].cells[idx].paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    set_run(p.add_run(text), size=8.5, bold=True, color=WHITE)
set_repeat_table_header(sequence.rows[0])
seq_rows = [
    ("1", "P0", "Stage 7M baseline", "Build + authenticated API smoke", "Baseline accepted"),
    ("2", "8A", "P0", "Visual audit + responsive screenshots", "v4.0 UI accepted"),
    ("3", "8B", "8A patterns", "Role and permission matrix", "v4.1 access accepted"),
    ("4", "8C", "P0 accounting/stock", "Return, GST and vendor reconciliation", "Return lifecycle accepted"),
    ("5", "8D", "8C movement rules", "Stock and valuation reconciliation", "Inventory authority accepted"),
    ("6", "8E", "8C/8D postings", "Journal, payment and banking reconciliation", "Financial controls accepted"),
    ("7", "8F", "Stable business rules", "Automated suites + restore drill", "Release candidate accepted"),
    ("8", "8G", "All prior stages", "Real integrations + go-live checklist", "Production approved"),
]
for row_data in seq_rows:
    cells = sequence.add_row().cells
    for idx, text in enumerate(row_data):
        if idx == 0:
            shade(cells[idx], CYAN)
        p = cells[idx].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        set_run(p.add_run(text), size=8.1, bold=(idx == 0), color=NAVY if idx == 0 else DARK)

add_text(doc, "Stage note discipline", style="Heading 2")
add_text(
    doc,
    "Each implementation package will add its plan to docs/planning, its durable technical notes to docs/modules, and its release notes plus validation evidence to the matching docs/stages folder. Issues discovered during development will be recorded in docs/planning/ISSUES.md before being marked fixed.",
)
add_text(doc, "Definition of done", style="Heading 2")
add_work_table(
    doc,
    [
        ("Code", "Backend, frontend, migrations and configuration are complete and scoped to the stage.", "Peer-readable diff and no unrelated refactor."),
        ("Build", "Release backend and production Nuxt builds complete.", "Build log recorded."),
        ("Runtime", "Docker services are healthy and the changed workflow is exercised.", "Smoke evidence recorded."),
        ("Data", "Migrations and older-volume repair paths are validated where applicable.", "Fresh and upgrade paths pass."),
        ("Security", "Role, scope and action permissions are checked.", "Unauthorized API and route attempts fail."),
        ("Business", "Stock, GST, payroll or accounting effects reconcile as applicable.", "Expected totals and ledger links match."),
        ("UX", "Desktop and mobile layouts are visually inspected.", "No overlap, clipping or unusable form."),
        ("Documentation", "Roadmap, TODO, issue register, stage notes and validation log are updated.", "Canonical docs remain current."),
    ],
)

add_page_break(doc)

# Risks and source appendix
add_text(doc, "3. Principal Risks and Controls", style="Heading 1")
risks = doc.add_table(rows=1, cols=3)
risks.style = "Table Grid"
set_table_widths(risks, [2700, 3300, 3360])
for idx, text in enumerate(("Risk", "Impact", "Control")):
    shade(risks.rows[0].cells[idx], NAVY)
    p = risks.rows[0].cells[idx].paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    set_run(p.add_run(text), size=9, bold=True, color=WHITE)
set_repeat_table_header(risks.rows[0])
risk_rows = [
    ("UI migration causes workflow regression", "Daily billing or purchase work slows or breaks", "Module-by-module rollout, existing routes preserved, legacy shell retained until acceptance"),
    ("Return documents diverge from stock/GST/accounting", "Incorrect vendor balance or tax reporting", "Single transactional posting service plus reconciliation tests"),
    ("Movement-ledger migration changes displayed stock", "Unexpected stock variance", "Pre-migration comparison report and controlled repair path"),
    ("Permission mismatch between client and API", "Unauthorized access or false denial", "Server-first policies plus automated role matrix"),
    ("External integration credentials are unavailable", "Stage cannot be production-validated", "Keep provider adapters configurable and mark credential-bound gates explicitly"),
    ("Automated tests arrive too late", "Hardening changes create regressions", "Add tests alongside each Stage 8C-E business rule, then consolidate in 8F"),
]
for risk, impact, control in risk_rows:
    cells = risks.add_row().cells
    for idx, text in enumerate((risk, impact, control)):
        p = cells[idx].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        set_run(p.add_run(text), size=8.4, bold=(idx == 0), color=RED if idx == 0 else DARK)

add_text(doc, "Source and provenance", style="Heading 1")
add_text(
    doc,
    "This roadmap was compiled from the Stage 7M repository baseline, docs/planning/CURRENT-ROADMAP.md, docs/planning/MASTER-TODO.md, docs/planning/ISSUES.md, docs/planning/TODO-RECONCILIATION.md, Stage 7 implementation notes, local build/runtime audit results, and the Stage 8A/8B direction retained in the locally preserved web-chat planning history.",
)
add_text(
    doc,
    "The private ChatGPT conversation named 'Garmetix Web 2' was not available in the local workspace as an export and was not accessible through a public shared link. Consequently, this document includes all available locally preserved Stage 8 content but does not claim verbatim or complete access to that private conversation.",
)
add_callout(
    doc,
    "Recommended first implementation package",
    "Priority 0 Baseline Stabilization: fix the empty dashboard response, remove compiler warnings, stabilize Nuxt font/build behavior, add aggregate validation, and run the authenticated Docker smoke test. This establishes a trustworthy base before v4.0 UI work.",
    fill=CYAN,
)

OUT.mkdir(parents=True, exist_ok=True)
doc.save(DOCX)
print(DOCX)
